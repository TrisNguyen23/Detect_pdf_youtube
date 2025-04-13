import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import requests
from bs4 import BeautifulSoup
from docx import Document
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer

def fetch_content():
    url = url_entry.get()
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        paragraphs = soup.find_all('p')
        content = '\n'.join([para.get_text() for para in paragraphs])
        content_text.delete('1.0', tk.END)
        content_text.insert(tk.END, content)
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể lấy nội dung từ URL: {e}")

def summarize_content():
    text = content_text.get('1.0', tk.END)
    if not text.strip():
        messagebox.showwarning("Cảnh báo", "Nội dung trống.")
        return
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count=3)
    summary_text.delete('1.0', tk.END)
    for sentence in summary:
        summary_text.insert(tk.END, str(sentence) + '\n')

def save_note():
    content = summary_text.get('1.0', tk.END)
    if not content.strip():
        messagebox.showwarning("Cảnh báo", "Nội dung ghi chú trống.")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc = Document()
        doc.add_heading('Ghi chú', 0)
        doc.add_paragraph(content)
        doc.save(file_path)
        messagebox.showinfo("Thành công", f"Ghi chú đã được lưu tại: {file_path}")

root = tk.Tk()
root.title("Ứng dụng Ghi chú và Tóm tắt Nội dung")
root.geometry("800x600")

url_label = tk.Label(root, text="Nhập URL:")
url_label.pack()
url_entry = tk.Entry(root, width=100)
url_entry.pack()

fetch_button = tk.Button(root, text="Lấy nội dung", command=fetch_content)
fetch_button.pack(pady=5)

content_label = tk.Label(root, text="Nội dung:")
content_label.pack()
content_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=100, height=10)
content_text.pack(pady=5)

summarize_button = tk.Button(root, text="Tóm tắt nội dung", command=summarize_content)
summarize_button.pack(pady=5)

summary_label = tk.Label(root, text="Tóm tắt:")
summary_label.pack()
summary_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=100, height=10)
summary_text.pack(pady=5)

save_button = tk.Button(root, text="Lưu ghi chú", command=save_note)
save_button.pack(pady=10)

root.mainloop()
